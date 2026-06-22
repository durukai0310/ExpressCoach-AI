import sys, os, re, traceback
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def escape_for_xml(text):
    """移除控制字符"""
    if not text:
        return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)

def process_inline(paragraph, text):
    """将带有 markdown 内联格式的文本添加到段落，保留粗体"""
    text = escape_for_xml(text)
    # 使用正则分割 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'

def add_rich_paragraph(doc, text, style_name='Normal', indent_left=None, italic=False):
    """添加带格式的段落"""
    if not text or not text.strip():
        return None
    p = doc.add_paragraph()
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    process_inline(p, text.strip())
    return p

def md_to_docx(md_path, docx_path):
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # 修改标题样式
    for i in [1, 2, 3, 4]:
        hs = doc.styles[f'Heading {i}']
        hs.font.name = '微软雅黑'
        hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    table_data = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip('\n').rstrip('\r')

        # 代码块
        if stripped.strip().startswith('```'):
            if in_code_block:
                # 输出代码块
                if code_lines:
                    for cl in code_lines:
                        p = doc.add_paragraph()
                        run = p.add_run(cl)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                        p.paragraph_format.left_indent = Cm(1)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                    code_lines = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            code_lines.append(stripped)
            i += 1
            continue

        # 空行
        if not stripped.strip():
            i += 1
            continue

        s = stripped.strip()

        # 分隔线
        if s.startswith('---') or s == '===' or all(c in '-＝=' for c in s):
            p = doc.add_paragraph()
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            run.font.size = Pt(8)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # H1
        if s.startswith('# ') and not s.startswith('## '):
            doc.add_heading(s[2:], level=1)
            i += 1
            continue

        # H2
        if s.startswith('## ') and not s.startswith('### '):
            doc.add_heading(s[3:], level=2)
            i += 1
            continue

        # H3
        if s.startswith('### ') and not s.startswith('#### '):
            doc.add_heading(s[4:], level=3)
            i += 1
            continue

        # H4
        if s.startswith('#### '):
            doc.add_heading(s[5:], level=4)
            i += 1
            continue

        # 表格
        if '|' in s:
            cells = [c.strip() for c in s.split('|')]
            # 去掉首尾空元素
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]

            # 检测是否分隔行 (--- 或 :---:)
            is_sep = all(re.match(r'^[-:]+$', c) for c in cells)
            if is_sep:
                i += 1
                continue

            table_data.append(cells)

            # 看下一行
            next_is_table = False
            if i + 1 < len(lines):
                ns = lines[i + 1].strip()
                if '|' in ns:
                    ncells = [c.strip() for c in ns.split('|')]
                    if ncells and ncells[0] == '':
                        ncells = ncells[1:]
                    if ncells and ncells[-1] == '':
                        ncells = ncells[:-1]
                    if not all(re.match(r'^[-:]+$', c) for c in ncells):
                        next_is_table = True

            if not next_is_table and table_data:
                # 输出表格
                num_cols = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Light Grid Accent 1'
                for ri, row in enumerate(table_data):
                    for ci in range(min(len(row), num_cols)):
                        cell = table.cell(ri, ci)
                        cell.text = escape_for_xml(row[ci])
                        for para in cell.paragraphs:
                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.space_before = Pt(0)
                            for run in para.runs:
                                run.font.size = Pt(9)
                                run.font.name = '微软雅黑'
                                if ri == 0:
                                    run.font.bold = True
                doc.add_paragraph()  # 表后空行
                table_data = []
            elif not next_is_table:
                table_data = []

            i += 1
            continue

        # 引用
        if s.startswith('> '):
            text = s[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'
            i += 1
            continue

        # 无序列表
        if re.match(r'^[-*]\s', s):
            text = re.sub(r'^[-*]\s+', '', s)
            p = doc.add_paragraph(style='List Bullet')
            # 清除默认文本
            p.clear()
            process_inline(p, text)
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+[\.\、]\s', s):
            text = re.sub(r'^\d+[\.\、]\s+', '', s)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            process_inline(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        process_inline(p, s)
        i += 1

    doc.save(docx_path)
    return True

# 主程序
desktop = r"C:\Users\chenyuxuan\Desktop"
project_docs = os.path.join(desktop, "expresscoach", "docs")

files = [
    (os.path.join(project_docs, "W4_详细任务分工_完整版.md"), os.path.join(desktop, "W4_详细任务分工_完整版.docx")),
    (os.path.join(desktop, "W4_成员A_详细操作手册.md"), os.path.join(desktop, "W4_成员A_详细操作手册.docx")),
    (os.path.join(desktop, "W4_成员B_详细操作手册.md"), os.path.join(desktop, "W4_成员B_详细操作手册.docx")),
    (os.path.join(desktop, "W4_成员C_详细操作手册.md"), os.path.join(desktop, "W4_成员C_详细操作手册.docx")),
]

for md_path, docx_path in files:
    try:
        print(f"转换: {os.path.basename(md_path)} -> {os.path.basename(docx_path)}")
        md_to_docx(md_path, docx_path)
        size_kb = os.path.getsize(docx_path) / 1024
        print(f"  OK ({size_kb:.1f} KB)")
    except Exception as e:
        print(f"  FAIL: {e}")

print("\n全部完成!")
